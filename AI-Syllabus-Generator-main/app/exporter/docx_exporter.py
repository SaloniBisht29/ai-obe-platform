import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from app.schemas.models import SyllabusResponse

EXPORTS_DIR = "outputs/exports"

def sf(run, size=11, bold=False, color=None, italic=False, font="Times New Roman"):
    run.font.name   = font
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def remove_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc  = cell._tc.get_or_add_tcPr()
            bdr = OxmlElement("w:tcBorders")
            for side in ["top","left","bottom","right","insideH","insideV"]:
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"),   "none")
                b.set(qn("w:sz"),    "0")
                b.set(qn("w:space"), "0")
                b.set(qn("w:color"), "auto")
                bdr.append(b)
            tc.append(bdr)

def shade_cell(cell, fill_hex):
    tc  = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tc.append(shd)

def red_footer(doc, programme, branch):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    txt = f"  Syllabus of {programme} \u2013 {branch}{'  '*25}PAGE 1"
    r   = p.add_run(txt)
    sf(r, size=9, bold=True, color=(255,255,255), font="Arial")
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "CC0000")
    pPr.append(shd)

def add_matrix_table(doc, matrix, col_headers, row_label, title):
    if not matrix:
        return
    p = doc.add_paragraph()
    sf(p.add_run(title), size=11, bold=True)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)

    rows = list(matrix.keys())
    cols = col_headers
    tbl  = doc.add_table(rows=len(rows)+1, cols=len(cols)+1)
    tbl.style = "Table Grid"

    # Header
    shade_cell(tbl.rows[0].cells[0], "1F4E79")
    r = tbl.rows[0].cells[0].paragraphs[0].add_run(row_label)
    sf(r, size=9, bold=True, color=(255,255,255))
    for j, col in enumerate(cols):
        cell = tbl.rows[0].cells[j+1]
        shade_cell(cell, "1F4E79")
        r = cell.paragraphs[0].add_run(col)
        sf(r, size=9, bold=True, color=(255,255,255))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data
    for i, row_key in enumerate(rows):
        row = tbl.rows[i+1]
        shade_cell(row.cells[0], "DCE6F1")
        r = row.cells[0].paragraphs[0].add_run(row_key)
        sf(r, size=9, bold=True)
        for j, col in enumerate(cols):
            val  = matrix.get(row_key, {}).get(col, 0)
            cell = row.cells[j+1]
            if   val == 3: shade_cell(cell, "C6EFCE"); display = "3"
            elif val == 2: shade_cell(cell, "FFEB9C"); display = "2"
            elif val == 1: shade_cell(cell, "FFCCCC"); display = "1"
            else:          display = "-"
            r2 = cell.paragraphs[0].add_run(display)
            sf(r2, size=9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    leg = doc.add_paragraph()
    sf(leg.add_run("Legend: "), size=9, bold=True)
    sf(leg.add_run("3=High  "), size=9, color=(0,100,0))
    sf(leg.add_run("2=Medium  "), size=9, color=(150,100,0))
    sf(leg.add_run("1=Low  "), size=9, color=(150,0,0))
    sf(leg.add_run("-=No Correlation"), size=9)
    leg.paragraph_format.space_after = Pt(8)


def export_syllabus_to_docx(syllabus: SyllabusResponse) -> str:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # ── University name ──
    univ = syllabus.university_name or "G.B. Pant Institute of Engineering & Technology, Pauri Garhwal"
    up   = doc.add_paragraph()
    up.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sf(up.add_run(univ.upper()), size=12, bold=True, color=(31,78,121))
    up.paragraph_format.space_after = Pt(4)

    # ── Syllabus heading ──
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sf(sp.add_run("Syllabus"), size=14, bold=True)
    sp.paragraph_format.space_after = Pt(6)

    # ── Course title ──
    title = syllabus.course_name
    if syllabus.course_code:
        title += f" ({syllabus.course_code})"
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sf(tp.add_run(title), size=13, bold=True)
    tp.paragraph_format.space_after = Pt(8)

    # ── Programme info ──
    parts = []
    if syllabus.programme:       parts.append(syllabus.programme.upper())
    if syllabus.education_level: parts.append(syllabus.education_level.title())
    if syllabus.year_of_study:   parts.append(f"Year {syllabus.year_of_study}")
    if syllabus.semester:        parts.append(f"Semester-{syllabus.semester}")
    if syllabus.branch:          parts.append(f"Branch: {syllabus.branch}")
    if parts:
        pi = doc.add_paragraph()
        sf(pi.add_run(" | ".join(parts)), size=10, italic=True, color=(80,80,80))
        pi.paragraph_format.space_after = Pt(4)

    # ── Standards ──
    if syllabus.standards:
        st = doc.add_paragraph()
        sf(st.add_run("Standards: "), size=10, bold=True, color=(31,78,121))
        sf(st.add_run(syllabus.standards), size=10, italic=True, color=(80,80,80))
        st.paragraph_format.space_after = Pt(4)

    # ── L:T:P + Credits ──
    ltp_tbl = doc.add_table(rows=1, cols=2)
    ltp_tbl.columns[0].width = Inches(3.0)
    ltp_tbl.columns[1].width = Inches(3.5)
    remove_borders(ltp_tbl)
    ll = ltp_tbl.cell(0,0).paragraphs[0]
    lr = ltp_tbl.cell(0,1).paragraphs[0]
    lr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sf(ll.add_run(f"L:T:P:: {syllabus.ltp or '3:1:0'}"), size=11, bold=True)
    sf(lr.add_run(f"Credits-{syllabus.credits or 4}"), size=11, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Hours and lectures ──
    if syllabus.total_hours:
        hp = doc.add_paragraph()
        sf(hp.add_run(f"Total Contact Hours: {syllabus.total_hours} hrs"), size=10, bold=True)
        if syllabus.total_lectures:
            hp.add_run("  |  ")
            sf(hp.add_run(f"Total Lectures: {syllabus.total_lectures}"), size=10, bold=True)
        hp.paragraph_format.space_after = Pt(6)

    # ── Exam pattern table ──
    if syllabus.exam_pattern:
        ep = doc.add_paragraph()
        sf(ep.add_run("ASSESSMENT PATTERN:"), size=11, bold=True)
        ep.paragraph_format.space_before = Pt(6)
        ep.paragraph_format.space_after  = Pt(4)

        exam_tbl = doc.add_table(rows=6, cols=2)
        exam_tbl.style = "Table Grid"
        exam_rows = [
            ("Assessment Component", "Marks"),
            ("Unit Test I", "10"),
            ("Unit Test II", "10"),
            ("Assignments", "5"),
            ("Attendance", "5"),
            ("End Semester Examination", "70"),
        ]
        for i, (label, marks) in enumerate(exam_rows):
            row = exam_tbl.rows[i]
            if i == 0:
                shade_cell(row.cells[0], "1F4E79")
                shade_cell(row.cells[1], "1F4E79")
                sf(row.cells[0].paragraphs[0].add_run(label), size=10, bold=True, color=(255,255,255))
                sf(row.cells[1].paragraphs[0].add_run(marks), size=10, bold=True, color=(255,255,255))
            elif i == 5:
                shade_cell(row.cells[0], "DCE6F1")
                shade_cell(row.cells[1], "DCE6F1")
                sf(row.cells[0].paragraphs[0].add_run(label), size=10, bold=True)
                sf(row.cells[1].paragraphs[0].add_run(marks), size=10, bold=True)
            else:
                sf(row.cells[0].paragraphs[0].add_run(label), size=10)
                sf(row.cells[1].paragraphs[0].add_run(marks), size=10)
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── Course Objectives ──
    if syllabus.course_objectives:
        p = doc.add_paragraph()
        sf(p.add_run("COURSE OBJECTIVES:"), size=11, bold=True)
        sf(p.add_run(" The objectives of this course are to:"), size=11)
        p.paragraph_format.space_after = Pt(4)
        for obj in syllabus.course_objectives:
            bp = doc.add_paragraph(style="List Number")
            sf(bp.add_run(obj), size=11)
            bp.paragraph_format.left_indent  = Inches(0.5)
            bp.paragraph_format.space_before = Pt(2)
            bp.paragraph_format.space_after  = Pt(2)
        doc.add_paragraph()

    # ── Course Outcomes — Full Indian NBA format ──
    if syllabus.course_outcomes:
        p = doc.add_paragraph()
        sf(p.add_run("COURSE OUTCOMES (COs):"), size=11, bold=True)
        p.paragraph_format.space_after = Pt(4)
        p2 = doc.add_paragraph()
        sf(p2.add_run("At the end of this course, the students will be able to:"), size=11)
        p2.paragraph_format.space_after = Pt(4)

        for co in syllabus.course_outcomes:
            bp = doc.add_paragraph()
            bp.paragraph_format.left_indent  = Inches(0.2)
            bp.paragraph_format.space_before = Pt(6)
            bp.paragraph_format.space_after  = Pt(2)
            sf(bp.add_run(f"{co.co_id}: "), size=11, bold=True, color=(31,78,121))
            sf(bp.add_run(co.text), size=11)

            # Bloom
            bl = doc.add_paragraph()
            bl.paragraph_format.left_indent = Inches(0.5)
            bl.paragraph_format.space_after = Pt(1)
            sf(bl.add_run("Bloom's Level: "), size=10, bold=True)
            sf(bl.add_run(f"{co.bloom_level} ({co.bloom_level_number}) — Verb: {co.bloom_verb}"), size=10)

            # PO mapping
            if co.mapped_pos:
                pm = doc.add_paragraph()
                pm.paragraph_format.left_indent = Inches(0.5)
                pm.paragraph_format.space_after = Pt(1)
                sf(pm.add_run("Mapped POs: "), size=10, bold=True)
                pos_str = ", ".join([f"{po} (Strength: {co.po_correlation.get(po,'?')})" for po in co.mapped_pos])
                sf(pm.add_run(pos_str), size=10)

            # PSO mapping
            if co.mapped_psos:
                psm = doc.add_paragraph()
                psm.paragraph_format.left_indent = Inches(0.5)
                psm.paragraph_format.space_after = Pt(1)
                sf(psm.add_run("Mapped PSOs: "), size=10, bold=True)
                pso_str = ", ".join([f"{pso} (Strength: {co.pso_correlation.get(pso,'?')})" for pso in co.mapped_psos])
                sf(psm.add_run(pso_str), size=10)

            # Attainment
            at = doc.add_paragraph()
            at.paragraph_format.left_indent = Inches(0.5)
            at.paragraph_format.space_after = Pt(1)
            sf(at.add_run("Attainment Target: "), size=10, bold=True)
            sf(at.add_run(f"{co.attainment_target} (Level {co.attainment_level})"), size=10)

            # Direct assessment
            if co.direct_assessment:
                da = doc.add_paragraph()
                da.paragraph_format.left_indent = Inches(0.5)
                da.paragraph_format.space_after = Pt(1)
                sf(da.add_run("Direct Assessment: "), size=10, bold=True)
                sf(da.add_run(", ".join(co.direct_assessment)), size=10)

            # Indirect assessment
            if co.indirect_assessment:
                ia = doc.add_paragraph()
                ia.paragraph_format.left_indent = Inches(0.5)
                ia.paragraph_format.space_after = Pt(4)
                sf(ia.add_run("Indirect Assessment: "), size=10, bold=True)
                sf(ia.add_run(", ".join(co.indirect_assessment)), size=10)

        doc.add_paragraph()

    # ── Attainment formulas ──
    af = doc.add_paragraph()
    sf(af.add_run("Attainment Formulas: "), size=10, bold=True, color=(31,78,121))
    af.paragraph_format.space_after = Pt(2)

    af2 = doc.add_paragraph()
    af2.paragraph_format.left_indent = Inches(0.3)
    sf(af2.add_run(f"CO: {syllabus.attainment_formula}"), size=10, italic=True)
    af2.paragraph_format.space_after = Pt(2)

    af3 = doc.add_paragraph()
    af3.paragraph_format.left_indent = Inches(0.3)
    sf(af3.add_run(f"PO: {syllabus.po_attainment_formula}"), size=10, italic=True)
    af3.paragraph_format.space_after = Pt(4)

    # ── Attainment levels table ──
    if syllabus.attainment_levels:
        alp = doc.add_paragraph()
        sf(alp.add_run("NBA ATTAINMENT LEVELS:"), size=10, bold=True)
        alp.paragraph_format.space_after = Pt(4)

        al_tbl = doc.add_table(rows=5, cols=2)
        al_tbl.style = "Table Grid"
        al_rows = [("Level", "Criteria")] + list(syllabus.attainment_levels.items())
        for i, (level, criteria) in enumerate(al_rows):
            row = al_tbl.rows[i]
            if i == 0:
                shade_cell(row.cells[0], "1F4E79")
                shade_cell(row.cells[1], "1F4E79")
                sf(row.cells[0].paragraphs[0].add_run(level),    size=9, bold=True, color=(255,255,255))
                sf(row.cells[1].paragraphs[0].add_run(criteria), size=9, bold=True, color=(255,255,255))
            else:
                sf(row.cells[0].paragraphs[0].add_run(level),    size=9, bold=True)
                sf(row.cells[1].paragraphs[0].add_run(criteria), size=9)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── CO-PO Matrix ──
    if syllabus.co_po_matrix:
        pos = ["PO1","PO2","PO3","PO4","PO5","PO6","PO7","PO8","PO9","PO10","PO11","PO12"]
        add_matrix_table(doc, syllabus.co_po_matrix, pos, "CO/PO",
                         "CO-PO MAPPING MATRIX (as per NBA SAR Table 1.4.2.1):")

    # ── CO-PSO Matrix ──
    if syllabus.co_pso_matrix:
        psos = ["PSO1","PSO2","PSO3"]
        add_matrix_table(doc, syllabus.co_pso_matrix, psos, "CO/PSO",
                         "CO-PSO MAPPING MATRIX:")

    # ── Units ──
    for unit in syllabus.units:
        hours = unit.hours or 8

        up = doc.add_paragraph()
        up.paragraph_format.space_before = Pt(10)
        up.paragraph_format.space_after  = Pt(4)
        sf(up.add_run(f"{unit.unit_id}: {unit.unit_title}:"), size=11, bold=True)
        up.add_run("\t\t\t\t")
        sf(up.add_run(f"({hours} hours)"), size=11, bold=True)

        # Satisfied COs
        if unit.satisfied_cos:
            sc = doc.add_paragraph()
            sf(sc.add_run("COs Satisfied: "), size=10, bold=True, color=(31,78,121))
            sf(sc.add_run(", ".join(unit.satisfied_cos)), size=10)
            sc.paragraph_format.space_after = Pt(2)

        # Lecture plan
        if unit.lecture_plan:
            lp = doc.add_paragraph()
            sf(lp.add_run("Lecture Plan: "), size=10, bold=True)
            sf(lp.add_run(unit.lecture_plan), size=10, italic=True)
            lp.paragraph_format.space_after = Pt(4)

        # Topics paragraph
        if unit.topics_paragraph:
            tp = doc.add_paragraph()
            sf(tp.add_run(unit.topics_paragraph), size=11)
            tp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            tp.paragraph_format.space_after = Pt(6)
        elif unit.topics:
            tp = doc.add_paragraph()
            sf(tp.add_run(", ".join(unit.topics) + f". ({hours} hours)"), size=11)
            tp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            tp.paragraph_format.space_after = Pt(6)

        # Topics bullets
        if unit.topics:
            p = doc.add_paragraph()
            sf(p.add_run("Topics Covered:"), size=11, bold=True)
            p.paragraph_format.space_after = Pt(2)
            for topic in unit.topics:
                bp = doc.add_paragraph(style="List Bullet")
                sf(bp.add_run(topic), size=11)
                bp.paragraph_format.left_indent  = Inches(0.5)
                bp.paragraph_format.space_before = Pt(1)
                bp.paragraph_format.space_after  = Pt(1)

        # CSOs
        if unit.unit_objectives:
            p = doc.add_paragraph()
            sf(p.add_run("Course Specific Objectives (CSOs):"), size=11, bold=True)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(2)
            for obj in unit.unit_objectives:
                bp = doc.add_paragraph(style="List Bullet")
                sf(bp.add_run(obj), size=11)
                bp.paragraph_format.left_indent  = Inches(0.5)
                bp.paragraph_format.space_before = Pt(1)
                bp.paragraph_format.space_after  = Pt(1)

        # Unit Outcomes
        if unit.unit_outcomes:
            p = doc.add_paragraph()
            sf(p.add_run("Course Specific Outcomes:"), size=11, bold=True)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(2)
            for out in unit.unit_outcomes:
                bp = doc.add_paragraph(style="List Bullet")
                sf(bp.add_run(out), size=11)
                bp.paragraph_format.left_indent  = Inches(0.5)
                bp.paragraph_format.space_before = Pt(1)
                bp.paragraph_format.space_after  = Pt(1)

        # Assessments
        if unit.assessments:
            p = doc.add_paragraph()
            sf(p.add_run("Assessments:"), size=11, bold=True)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(2)
            for a in unit.assessments:
                bp = doc.add_paragraph(style="List Bullet")
                sf(bp.add_run(a), size=11)
                bp.paragraph_format.left_indent  = Inches(0.5)
                bp.paragraph_format.space_before = Pt(1)
                bp.paragraph_format.space_after  = Pt(1)

        # Readings
        if unit.readings:
            p = doc.add_paragraph()
            sf(p.add_run("Readings:"), size=11, bold=True)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(2)
            for rd in unit.readings:
                bp = doc.add_paragraph(style="List Bullet")
                sf(bp.add_run(rd), size=11)
                bp.paragraph_format.left_indent  = Inches(0.5)
                bp.paragraph_format.space_before = Pt(1)
                bp.paragraph_format.space_after  = Pt(1)

        doc.add_paragraph()

    # ── CQI Plan ──
    if syllabus.cqi_plan:
        p = doc.add_paragraph()
        sf(p.add_run("CONTINUOUS QUALITY IMPROVEMENT (CQI) PLAN:"), size=11, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(4)
        cp = doc.add_paragraph()
        sf(cp.add_run(syllabus.cqi_plan), size=11)
        cp.paragraph_format.space_after = Pt(6)

    # ── Lesson plan note ──
    if syllabus.lesson_plan_note:
        lp = doc.add_paragraph()
        sf(lp.add_run("Lesson Plan: "), size=10, bold=True, color=(31,78,121))
        sf(lp.add_run(syllabus.lesson_plan_note), size=10, italic=True)
        lp.paragraph_format.space_after = Pt(6)

    # ── Textbooks ──
    if syllabus.textbooks:
        p = doc.add_paragraph()
        sf(p.add_run("BOOKS:"), size=11, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(6)
        for book in syllabus.textbooks:
            bp = doc.add_paragraph(style="List Number")
            sf(bp.add_run(book), size=11)
            bp.paragraph_format.left_indent  = Inches(0.5)
            bp.paragraph_format.space_before = Pt(2)
            bp.paragraph_format.space_after  = Pt(2)
        doc.add_paragraph()

    # ── YouTube ──
    if syllabus.youtube_resources:
        p = doc.add_paragraph()
        sf(p.add_run("YOUTUBE & VIDEO RESOURCES:"), size=11, bold=True)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(4)
        for res in syllabus.youtube_resources:
            bp = doc.add_paragraph(style="List Number")
            sf(bp.add_run(res), size=11)
            bp.paragraph_format.left_indent  = Inches(0.5)
            bp.paragraph_format.space_before = Pt(2)
            bp.paragraph_format.space_after  = Pt(2)
        doc.add_paragraph()

    # ── Open Source ──
    if syllabus.open_source_resources:
        p = doc.add_paragraph()
        sf(p.add_run("OPEN SOURCE & ONLINE RESOURCES (SWAYAM/NPTEL):"), size=11, bold=True)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(4)
        for res in syllabus.open_source_resources:
            bp = doc.add_paragraph(style="List Number")
            sf(bp.add_run(res), size=11)
            bp.paragraph_format.left_indent  = Inches(0.5)
            bp.paragraph_format.space_before = Pt(2)
            bp.paragraph_format.space_after  = Pt(2)

    # ── NAAC IQAC note ──
    if syllabus.naac_iqac_note:
        doc.add_paragraph()
        ni = doc.add_paragraph()
        sf(ni.add_run("Quality Compliance: "), size=10, bold=True, color=(31,78,121))
        sf(ni.add_run(syllabus.naac_iqac_note), size=10, italic=True, color=(80,80,80))
        ni.paragraph_format.space_after = Pt(8)

    # ── Red footer ──
    prog = syllabus.programme.upper() if syllabus.programme else "BTECH"
    br   = syllabus.branch or "Computer Science and Engineering"
    red_footer(doc, prog, br)

    # ── Save ──
    os.makedirs(EXPORTS_DIR, exist_ok=True)
    safe  = syllabus.course_name.replace(" ", "_")
    code  = f"_{syllabus.course_code}"       if syllabus.course_code   else ""
    prog  = syllabus.programme.upper()       if syllabus.programme     else "general"
    yr    = f"_Year{syllabus.year_of_study}" if syllabus.year_of_study else ""
    sem   = f"_S{syllabus.semester}"         if syllabus.semester      else ""
    fname = f"{EXPORTS_DIR}/{safe}{code}_{prog}{yr}{sem}_syllabus.docx"
    doc.save(fname)
    print(f"DOCX saved: {fname}")
    return os.path.abspath(fname)